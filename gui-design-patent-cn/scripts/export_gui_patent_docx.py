#!/usr/bin/env python3
"""
Assemble standard CN GUI filing Word `02界面说明.docx` + matching JPG bundle.

Designed for templates that match references/company-word-template-rules.md:
  page01 cover, page02 meta table, page03+ body starting at 「界面说明如下：」.

Depends on python-docx + Pillow (see ../requirements.txt).
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

ANCHOR_TEXT = "界面说明如下："
DEFAULT_PIC_WIDTH_IN = 5270500 / 914400  # matches typical template screenshots (~14.65 cm)


def _delete_paragraph(paragraph: Any) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is None:
        return
    parent.remove(el)


def truncate_body_after_anchor(doc: Document, anchor: str = ANCHOR_TEXT) -> int:
    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == anchor:
            anchor_idx = i
            break
    if anchor_idx is None:
        raise ValueError(f"模板中未找到正文锚点段落：{anchor!r}")
    while len(doc.paragraphs) > anchor_idx + 1:
        _delete_paragraph(doc.paragraphs[-1])
    return anchor_idx


def _set_cell_text(cell: Any, text: str) -> None:
    cell.text = text


def _coalesce(value: Any, fallback: str = "待补充") -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def normalize_inventors(case: dict[str, Any]) -> list[dict[str, str]]:
    inv = case.get("inventors")
    if isinstance(inv, list) and inv:
        out: list[dict[str, str]] = []
        for row in inv:
            if not isinstance(row, dict):
                continue
            name = _coalesce(row.get("name"), "")
            id_card = _coalesce(row.get("id_card"), "")
            if name == "" and id_card == "":
                continue
            out.append({"name": name, "id_card": id_card})
        if out:
            return out[:8]
    name = _coalesce(case.get("inventor"), "")
    if name == "":
        return []
    return [{"name": name, "id_card": _coalesce(case.get("id_card"), "")}]


def fill_meta_table(doc: Document, case: dict[str, Any]) -> None:
    if not doc.tables:
        raise ValueError("模板缺少信息表（期望 tables[0] 为元数据表格）")

    table = doc.tables[0]

    date_text = _coalesce(case.get("date"), "")
    if date_text == "待补充":
        date_text = datetime_today_cn()

    patent_type = _coalesce(case.get("patent_type"), "GUI外观设计专利")
    domain = _coalesce(case.get("domain"), "软件界面外观设计")
    product_name = _coalesce(case.get("product_name"), "")
    if product_name in ("", "待补充"):
        cands = case.get("product_name_candidates")
        if isinstance(cands, list) and cands:
            product_name = _coalesce(cands[0], "待补充")

    applicant = _coalesce(case.get("applicant"), "")
    phone = _coalesce(case.get("phone"), "")
    email = _coalesce(case.get("email"), "")
    fixed_phone = _coalesce(case.get("fixed_phone"), "(补充信息)")

    # Row 0
    _set_cell_text(table.rows[0].cells[0], date_text)
    _set_cell_text(table.rows[0].cells[5], patent_type)
    _set_cell_text(table.rows[0].cells[10], domain)

    # Row 1 patent name span starts ~cell 2
    _set_cell_text(table.rows[1].cells[2], product_name)

    # Applicant
    _set_cell_text(table.rows[2].cells[2], applicant)

    inventors = normalize_inventors(case)
    slot_pairs = [
        (4, "left"),
        (4, "right"),
        (5, "left"),
        (5, "right"),
        (6, "left"),
        (6, "right"),
        (7, "left"),
        (7, "right"),
    ]
    for idx, (row_idx, side) in enumerate(slot_pairs, start=1):
        row = table.rows[row_idx]
        if side == "left":
            _set_cell_text(row.cells[1], str(idx))
            name = inventors[idx - 1]["name"] if idx <= len(inventors) else ""
            id_card = inventors[idx - 1]["id_card"] if idx <= len(inventors) else ""
            _set_cell_text(row.cells[3], name)
            _set_cell_text(row.cells[4], id_card)
        else:
            seq = idx
            _set_cell_text(row.cells[6], str(seq))
            name = inventors[idx - 1]["name"] if idx <= len(inventors) else ""
            id_card = inventors[idx - 1]["id_card"] if idx <= len(inventors) else ""
            _set_cell_text(row.cells[8], name)
            _set_cell_text(row.cells[11], id_card)

    contact_phone = _coalesce(case.get("inventor_contact_phone"), phone)
    if contact_phone == "待补充":
        contact_phone = ""

    row8 = table.rows[8]
    _set_cell_text(row8.cells[0], f"发明联系方式：{contact_phone}")
    _set_cell_text(row8.cells[4], f"固定电话：{fixed_phone}")
    _set_cell_text(row8.cells[9], f"手机：{phone}")

    row9 = table.rows[9]
    _set_cell_text(row9.cells[0], f"邮箱：{email}")

def datetime_today_cn() -> str:
    now = datetime.now()
    return f"{now.year}年{now.month}月{now.day}日"


def build_summary_intro(case: dict[str, Any]) -> str:
    custom = case.get("summary_intro")
    if isinstance(custom, str) and custom.strip():
        return custom.strip()
    carrier = _coalesce(case.get("carrier_product"), "电子设备")
    filing_goal = _coalesce(case.get("filing_goal"), "对应业务用途")
    return f"本外观设计产品用于在{carrier}上展示，用于{filing_goal}的图形用户界面。"


def build_summary_bullets(case: dict[str, Any]) -> list[str]:
    bullets = case.get("summary_bullets")
    if isinstance(bullets, list) and bullets:
        return [str(x).strip() for x in bullets if str(x).strip()]
    states = case.get("interface_states")
    if isinstance(states, list) and states:
        names: list[str] = []
        for st in states:
            if isinstance(st, dict):
                names.append(_coalesce(st.get("name"), ""))
            else:
                names.append(str(st))
        return [n for n in names if n and n != "待补充"]
    return ["待补充状态列表"]


def build_design_points(case: dict[str, Any]) -> str:
    custom = case.get("design_points_sentence")
    if isinstance(custom, str) and custom.strip():
        return custom.strip()
    novelty = case.get("core_novelty")
    parts: list[str] = []
    if isinstance(novelty, list):
        parts = [str(x).strip() for x in novelty if str(x).strip()]
    joined = "；".join(parts) if parts else "待结合状态图补充客观的设计要点表述。"
    return f"本外观设计的设计要点在于{joined}。"


def append_standard_body(doc: Document, case: dict[str, Any], image_width_in: float) -> None:
    doc.add_paragraph()

    p_abs = doc.add_paragraph()
    r = p_abs.add_run("摘要：")
    r.bold = True

    intro_base = build_summary_intro(case).strip().rstrip("。")
    opening = doc.add_paragraph()
    opening.add_run(f"{intro_base}。包括：")

    bullets = build_summary_bullets(case)
    for item in bullets:
        bp = doc.add_paragraph(style="List Number")
        bp.add_run(item)

    dp = doc.add_paragraph()
    dp.add_run(build_design_points(case))

    sec = doc.add_paragraph()
    r2 = sec.add_run("各界面状态说明：")
    r2.bold = True

    states = case.get("interface_states")
    if not isinstance(states, list) or not states:
        raise ValueError("interface_states 不能为空（导出 Word 需要逐张状态图）")

    base_dir = Path(case.get("_resolved_json_dir") or ".")

    for i, st in enumerate(states, start=1):
        if not isinstance(st, dict):
            raise ValueError("interface_states 的每项必须是对象")
        name = _coalesce(st.get("name"), f"状态{i}")
        desc = st.get("description") or st.get("summary") or ""
        desc = str(desc).strip()
        if not desc:
            desc = "请根据该状态截图补充客观界面描述（禁止主观评价词）。"

        title = doc.add_paragraph()
        rt = title.add_run(f"{name}（变化状态图{i:02d}）")
        rt.bold = True

        body_p = doc.add_paragraph()
        body_p.add_run(desc)

        img_rel = st.get("image_path")
        if not img_rel:
            raise ValueError(f"状态「{name}」缺少 image_path（必须是可直接读取的单张截图路径）")

        img_path = (base_dir / str(img_rel)).resolve()
        if not img_path.is_file():
            raise FileNotFoundError(f"找不到图片文件：{img_path}")

        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_pic = pic_p.add_run()
        run_pic.add_picture(str(img_path), width=Inches(image_width_in))

        doc.add_paragraph()

    inter_title = doc.add_paragraph()
    r3 = inter_title.add_run("交互说明：")
    r3.bold = True

    inter = case.get("interaction_summary")
    if isinstance(inter, str) and inter.strip():
        inter_text = inter.strip()
    else:
        inter_text = _default_interaction(len(states))

    # 按换行拆段渲染，符合「切换方式 / 保护要点 / 不涉及部分」分行显示惯例
    for chunk in inter_text.split("\n"):
        chunk = chunk.strip()
        if chunk:
            doc.add_paragraph(chunk)


def _default_interaction(state_count: int) -> str:
    """兜底交互说明 — 必须包含三层结构（切换方式 + 保护要点 + 不涉及部分）。

    详见 references/drafting-workflow.md Step F。即使内容粗糙，也要把"保护要点"和
    "不涉及部分"两个白/黑名单层写出来，避免 GUI 外观设计专利保护范围被错误理解为
    "文字内容 + LOGO + 图片"层面。建议每个 case.json 都显式提供 interaction_summary，
    本兜底仅用于极简情况。
    """
    if state_count <= 1:
        return "请根据最终保留的状态图补充交互说明（必须包含三层：切换方式、保护要点、不涉及部分）。"
    last_idx = f"变化状态图{state_count:02d}"
    return (
        f"各界面通过用户操作依次展示或动态切换，上述界面按变化状态图01-{last_idx}状态分别展示。"
        "保护要点在于各状态共同构成的整体视觉编排结构与模块组合关系。\n"
        "不涉及的部分：文字内容本身、品牌LOGO本身、图片素材本身、具体数据字段内容、"
        "网页背景色值、动画效果。"
    )


def sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|\s]+', "_", name.strip())
    name = name.strip("._") or "state"
    return name[:120]


def export_jpg_bundle(case: dict[str, Any], out_dir: Path, quality: int = 92) -> list[Path]:
    try:
        from PIL import Image
    except ImportError as exc:  # pragma: no cover - guarded by requirements
        raise RuntimeError("导出 JPG 需要安装 Pillow：pip install Pillow") from exc

    states = case.get("interface_states")
    if not isinstance(states, list):
        return []

    base_dir = Path(case.get("_resolved_json_dir") or ".")
    img_dir = out_dir / "jpg"
    img_dir.mkdir(parents=True, exist_ok=True)
    exported: list[Path] = []

    for i, st in enumerate(states, start=1):
        if not isinstance(st, dict):
            continue
        name = _coalesce(st.get("name"), f"状态{i}")
        rel = st.get("image_path")
        if not rel:
            continue
        src = (base_dir / str(rel)).resolve()
        if not src.is_file():
            raise FileNotFoundError(f"找不到图片文件：{src}")

        stem = sanitize_filename(name)
        dst = img_dir / f"变化状态图{i:02d}_{stem}.jpg"

        im = Image.open(src)
        rgb = im.convert("RGB")
        rgb.save(dst, format="JPEG", quality=quality, optimize=True)
        exported.append(dst)

    return exported


def resolve_case_paths(case_path: Path, case: dict[str, Any]) -> dict[str, Any]:
    case["_resolved_json_dir"] = str(case_path.parent.resolve())
    return case


def export_case(
    case_path: Path,
    template_path: Path,
    output_dir: Path,
    docx_name: str = "02界面说明.docx",
    image_width_in: float = DEFAULT_PIC_WIDTH_IN,
    skip_jpg: bool = False,
) -> tuple[Path, list[Path]]:
    case = json.loads(case_path.read_text(encoding="utf-8"))
    case = resolve_case_paths(case_path, case)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_docx = output_dir / docx_name
    shutil.copyfile(template_path, out_docx)

    doc = Document(str(out_docx))
    fill_meta_table(doc, case)
    truncate_body_after_anchor(doc)
    append_standard_body(doc, case, image_width_in=image_width_in)
    doc.save(str(out_docx))

    jpgs: list[Path] = []
    if not skip_jpg:
        jpgs = export_jpg_bundle(case, output_dir)

    return out_docx, jpgs


def parse_args(argv: list[str]) -> argparse.Namespace:
    p = argparse.ArgumentParser(description="导出 02界面说明.docx + JPG 图片包")
    p.add_argument("--input", required=True, help="结构化案件 JSON（UTF-8）")
    p.add_argument(
        "--template",
        default=None,
        help="Word 模板路径（默认：与本脚本相邻的 ../assets/word-template.docx）",
    )
    p.add_argument("--output-dir", required=True, help="输出目录（将生成 docx 与 jpg/）")
    p.add_argument("--docx-name", default="02界面说明.docx", help="输出 Word 文件名")
    p.add_argument(
        "--image-width-in",
        type=float,
        default=DEFAULT_PIC_WIDTH_IN,
        help="正文插图宽度（英寸），默认按模板常用截图宽度",
    )
    p.add_argument("--skip-jpg", action="store_true", help="仅导出 Word，不生成 JPG 目录")
    return p.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    case_path = Path(args.input).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()

    script_dir = Path(__file__).resolve().parent
    default_tpl = script_dir.parent / "assets" / "word-template.docx"
    template_path = Path(args.template).expanduser().resolve() if args.template else default_tpl

    if not case_path.is_file():
        print(f"找不到输入 JSON：{case_path}", file=sys.stderr)
        return 2
    if not template_path.is_file():
        print(f"找不到模板 DOCX：{template_path}", file=sys.stderr)
        return 2

    docx_out, jpgs = export_case(
        case_path,
        template_path,
        output_dir,
        docx_name=args.docx_name,
        image_width_in=args.image_width_in,
        skip_jpg=args.skip_jpg,
    )

    print(f"已生成 Word：{docx_out}")
    if jpgs:
        print(f"已导出 JPG {len(jpgs)} 张 -> {jpgs[0].parent}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
